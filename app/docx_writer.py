"""DOCX writer module for creating Study Context Guides."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re


class DOCXWriter:
    """Creates formatted DOCX Study Context Guides."""
    
    def create_study_guide_docx(
        self, 
        content: str, 
        slides_filename: str = "Lecture Deck",
        course_code: str = ""
    ) -> bytes:
        """
        Create a formatted DOCX from study guide content.
        
        Args:
            content: Generated study guide text (markdown-like format)
            slides_filename: Name of the slides file
            course_code: Course code
            
        Returns:
            DOCX file as bytes
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title_text = f"Study Context Guide — {slides_filename}"
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add course code if provided
        if course_code:
            course_para = doc.add_paragraph(f"Course: {course_code}")
            course_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            course_para.runs[0].font.size = Pt(12)
            course_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()  # Spacing
        
        # Parse and format content
        self._parse_and_format_content(doc, content)
        
        # Save to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
    
    def _parse_and_format_content(self, doc: Document, content: str):
        """
        Parse markdown-like content and format it in the document.
        
        Args:
            doc: Document object
            content: Text content to parse
        """
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Main headings (typically numbered or all caps)
            if self._is_main_heading(line):
                doc.add_heading(line.strip('#').strip(), level=1)
            
            # Sub-headings (bold text with ** or ##)
            elif line.strip().startswith('##') or (line.strip().startswith('**') and line.strip().endswith('**')):
                heading_text = line.strip('#').strip('*').strip()
                doc.add_heading(heading_text, level=2)
            
            # Bold text inline
            elif '**' in line and line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)
            
            # Bullet points
            elif line.strip().startswith(('-', '•', '*')) and len(line.strip()) > 2:
                text = line.strip().lstrip('-•*').strip()
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, text)
            
            # Regular paragraph
            elif line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)
            
            # Empty line - add spacing
            else:
                if i > 0 and lines[i-1].strip():  # Only add spacing if previous line had content
                    doc.add_paragraph()
            
            i += 1
    
    def _is_main_heading(self, line: str) -> bool:
        """Check if line is a main heading."""
        line = line.strip()
        
        # Numbered sections (1. CONCEPT MAP)
        if re.match(r'^\d+\.\s+[A-Z\s]+$', line):
            return True
        
        # All caps headings
        if line.isupper() and len(line.split()) >= 2 and len(line) < 60:
            return True
        
        # Single # heading
        if line.startswith('#') and not line.startswith('##'):
            return True
        
        return False
    
    def _add_formatted_text(self, paragraph, text: str):
        """
        Add text to paragraph with inline formatting (bold, etc.).
        
        Args:
            paragraph: Paragraph object
            text: Text with markdown-like formatting
        """
        # Split by ** for bold sections
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part.strip('*'))
                run.font.bold = True
            else:
                # Regular text
                paragraph.add_run(part)
